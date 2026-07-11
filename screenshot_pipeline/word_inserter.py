from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

from .models import ScreenshotConfig, ScreenshotTask


def _paragraphs(document):
    seen: set[int] = set()
    for paragraph in document.paragraphs:
        if id(paragraph._p) not in seen:
            seen.add(id(paragraph._p))
            yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if id(paragraph._p) not in seen:
                        seen.add(id(paragraph._p))
                        yield paragraph


def _find_anchor(document, anchor: str):
    for paragraph in _paragraphs(document):
        if anchor and anchor in paragraph.text:
            return paragraph
    return None


def _new_paragraph_after(paragraph) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    return Paragraph(element, paragraph._parent)


def _insert_images_after(paragraph, assets, page_break_between: bool, max_width_inches: float) -> None:
    cursor = paragraph
    for index, asset in enumerate(assets):
        if index and page_break_between:
            cursor = _new_paragraph_after(cursor)
            cursor.add_run().add_break()
        image_paragraph = _new_paragraph_after(cursor)
        image_paragraph.alignment = paragraph.alignment
        image_paragraph.add_run().add_picture(str(asset.stored_path), width=Inches(max_width_inches))
        cursor = image_paragraph


def insert_screenshots(
    template_path: str | Path,
    output_path: str | Path,
    task: ScreenshotTask,
    config: ScreenshotConfig,
    max_width_inches: float = 6.3,
) -> list[str]:
    document = Document(str(template_path))
    warnings: list[str] = []
    used_anchors: set[str] = set()
    for item in config.items:
        assets = sorted(task.assets.get(item.item_id, []), key=lambda asset: asset.order)
        if not assets:
            continue
        paragraph = _find_anchor(document, item.anchor)
        if paragraph is None:
            warnings.append(f"截图项目“{item.display_name}”未找到 Word 插入位置：{item.anchor}")
            continue
        _insert_images_after(paragraph, assets, item.page_break_between, max_width_inches)
        used_anchors.add(item.anchor)
    document.save(str(output_path))
    return warnings
