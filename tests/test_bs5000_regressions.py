from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from bs_section_filler.excel_blocks import _build_formula_evaluator
from bs_section_filler.utils import normalize_text, physical_table_matrix
from bs_section_filler.report_models import generate_report_for_model


ROOT = Path(__file__).resolve().parents[1]
EXCEL = ROOT / "BS5000" / "BS-5000 校准报告.xlsx"
TEMPLATE = ROOT / "BS5000" / "BS-5000系列_全自动生化分析仪仪器校准报告_V1.0_CH.docx"


def row_values(table, row_index: int) -> list[str]:
    return [" ".join(cell.text.split()) for cell in table.rows[row_index].cells]


class BS5000RegressionTests(unittest.TestCase):
    def test_ise_module1_carryover_formulas_are_recalculated(self) -> None:
        values = load_workbook(EXCEL, data_only=True)["1"]
        formulas = load_workbook(EXCEL, data_only=False)["1"]
        evaluator = _build_formula_evaluator(values, formulas)

        expected = {
            "C690": 0.006128481759268109,
            "F690": 0.00632661189903471,
            "I690": 0.008786713814806236,
            "C691": 0.003942335933081186,
            "F691": 0.005827427746712665,
            "I691": 0.004448102658629486,
        }
        for address, expected_value in expected.items():
            cell = formulas[address]
            self.assertAlmostEqual(evaluator(cell.row, cell.column), expected_value, places=7, msg=address)
        self.assertEqual(evaluator(formulas["C694"].row, formulas["C694"].column), "Pass")
        self.assertEqual(evaluator(formulas["F694"].row, formulas["F694"].column), "Pass")
        self.assertEqual(evaluator(formulas["I694"].row, formulas["I694"].column), "Pass")

    def _generate(self, temp_dir: str) -> Document:
        out = Path(temp_dir) / "bs5000.docx"
        generate_report_for_model(
            model_key="bs5000",
            template_path=TEMPLATE,
            excel_paths={"m1": EXCEL},
            output_path=out,
            formula_policy="all",
            overwrite_nonblank=True,
            delete_tail_sections=False,
        )
        return Document(out)

    def test_dark_current_inner_and_outer_regions_are_both_filled(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            doc = self._generate(temp_dir)
            table = doc.tables[6]

            self.assertEqual(row_values(table, 3)[1:], [
                "8", "8", "11", "8", "9", "9", "9", "10",
                "9", "7", "14", "8", "9", "11", "11", "10",
            ])
            self.assertEqual(row_values(table, 4)[1:], [
                "0.47", "0.5", "0.15", "0.12", "0.09", "0.09", "0.09", "0.09",
                "0.08", "0.08", "0.08", "0.07", "0.08", "0.08", "0.07", "0.08",
            ])
            self.assertEqual(row_values(table, 7)[1], "■ Pass")

            self.assertEqual(row_values(table, 11)[1:], [
                "5", "9", "9", "12", "11", "9", "10", "9",
                "10", "11", "12", "11", "6", "11", "14", "12",
            ])
            self.assertEqual(row_values(table, 12)[1:], [
                "0.49", "0.49", "0.15", "0.12", "0.1", "0.09", "0.1", "0.09",
                "0.08", "0.08", "0.07", "0.08", "0.07", "0.08", "0.08", "0.07",
            ])
            self.assertEqual(row_values(table, 15)[1], "■ Pass")

    def test_450nm_inner_max_absorbance_uses_the_excel_threshold_values(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            doc = self._generate(temp_dir)
            values = physical_table_matrix(doc.tables[14])[7]

            # Excel says "不小于2" while Word says "大于2".  Both describe
            # one threshold row and must use the same seven source results.
            self.assertEqual(values[5:], ["否", "否", "是", "是", "是", "是", "是"])

    def test_450nm_inner_relative_bias_uses_the_excel_result_row(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            doc = self._generate(temp_dir)
            values = physical_table_matrix(doc.tables[14])[8]

            self.assertEqual(
                values[5:],
                ["-0.08%", "-1.14%", "-2.14%", "-0.85%", "-1.11%", "-0.62%", "-1.26%"],
            )

    def test_sampling_error_uses_the_sampling_accuracy_indicator(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            doc = self._generate(temp_dir)

            # Excel calls this row "加样准确度指标" and the Word form calls it
            # "加样误差".  The displayed label is preserved, only its source is shared.
            for table_index in (18, 19):
                values = physical_table_matrix(doc.tables[table_index])[25]
                self.assertEqual(values[2::2], ["±5%", "±5%", "±5%"])

    def test_generated_word_uses_current_excel_values_for_ise_module1(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            doc = self._generate(temp_dir)
            table = doc.tables[48]
            row2 = row_values(table, 2)
            row6 = row_values(table, 6)
            row7 = row_values(table, 7)
            row9 = row_values(table, 9)

            self.assertEqual(row_values(table, 0)[1::3], ["Na", "K", "Cl"])
            self.assertEqual(row2[1:], ["98.58", "180.3", "99.11", "1.48", "7.54", "1.5", "76.99", "157.73", "77.28"])
            self.assertEqual(row6[1::3], ["0.61%", "0.63%", "0.88%"])
            self.assertEqual(row7[1::3], ["0.39%", "0.58%", "0.44%"])
            self.assertTrue(all(value == "■ Pass" for value in row9[1:]))

    def test_electrolyte_na_headers_are_not_treated_as_placeholders(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            doc = self._generate(temp_dir)
            cases = [
                (42, 1, ["项目", "NA", "K", "CL"]),
                (43, 1, ["项目次数", "NA", "K", "CL"]),
                (47, 0, ["项目测试时间", "NA", "K", "CL"]),
                (49, 1, ["项目", "NA", "K", "CL"]),
                (50, 1, ["项目次数", "NA", "K", "CL"]),
                (54, 0, ["项目测试时间", "NA", "K", "CL"]),
            ]

            for table_index, row_index, expected in cases:
                actual = [normalize_text(value) for value in row_values(doc.tables[table_index], row_index)[:4]]
                self.assertEqual(actual, expected, table_index)
            self.assertIn("NA线性范围验证", row_values(doc.tables[44], 0)[0])
            self.assertIn("NA线性范围验证", row_values(doc.tables[51], 0)[0])

    def test_module2_ise_empty_sources_do_not_emit_formula_or_template_results(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            doc = self._generate(temp_dir)

            accuracy = physical_table_matrix(doc.tables[49])
            self.assertTrue(all(value == "" for value in accuracy[10][1:]), accuracy[10])

            precision = physical_table_matrix(doc.tables[50])
            for row_index in (12, 13, 14, 16):
                self.assertTrue(all(value == "" for value in precision[row_index][1:]), (row_index, precision[row_index]))

            for table_index in (51, 52, 53):
                matrix = physical_table_matrix(doc.tables[table_index])
                self.assertEqual(matrix[8][0], "", (table_index, matrix[8]))
                self.assertTrue(all(value == "" for value in matrix[9][1:]), (table_index, matrix[9]))
                self.assertEqual(matrix[10][0], "", (table_index, matrix[10]))

            stability = physical_table_matrix(doc.tables[54])
            self.assertTrue(all(value == "" for value in stability[7][1:]), stability[7])
            self.assertTrue(all(value == "" for value in stability[9][1:]), stability[9])

            carryover = physical_table_matrix(doc.tables[55])
            for row_index in (6, 7, 9):
                self.assertTrue(all(value == "" for value in carryover[row_index][1:]), (row_index, carryover[row_index]))

    def test_sample_carryover_inner_outer_rows_are_filled_independently(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            doc = self._generate(temp_dir)
            table = doc.tables[36]
            inner_measurements = row_values(table, 2)[2:]
            outer_measurements = row_values(table, 8)[2:]

            self.assertEqual(inner_measurements, ["0.007141", "0.000161", "0.001680", "-0.000980", "0.000301"])
            self.assertEqual(outer_measurements, ["0.001740", "-0.000320", "0.002121", "-0.000172", "0.000428"])
            self.assertNotEqual(inner_measurements, outer_measurements)
            self.assertEqual(row_values(table, 3)[2], "0.84")
            self.assertEqual(row_values(table, 4)[2], "0.001661%")
            self.assertTrue(all(value == "■ Pass" for value in row_values(table, 6)[2:]))
            self.assertEqual(row_values(table, 9)[2], "0.9")
            self.assertEqual(row_values(table, 10)[2], "0.000759%")
            self.assertTrue(all(value == "■ Pass" for value in row_values(table, 12)[2:]))

    def test_clinical_precision_is_filled_but_accuracy_stays_blank_without_target_values(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            doc = self._generate(temp_dir)
            precision = doc.tables[40]
            accuracy = doc.tables[41]

            # 精密度有明确的原始测量数据，应正常填充。
            self.assertEqual(row_values(precision, 4)[1:8], ["60.1", "47.6", "168.5", "1.28", "6.73", "50.76", "2.17"])
            self.assertEqual(row_values(precision, 24)[1:8], ["60.456", "47.75", "168.34", "1.28", "6.7706", "50.6558", "2.19"])
            self.assertTrue(all(value == "■ Pass" for value in row_values(precision, 30)[1:8]))

            # 准确性表的“质控靶值”为空，低限/高限/是否在控都没有有效数据来源。
            # 除项目名及模板自带占位文字外，不允许写入单位、0、实测值或 Fail。
            for row_index in list(range(2, 11)) + list(range(13, 22)):
                values = row_values(accuracy, row_index)
                self.assertTrue(all(value in {"", "PassFail"} for value in values[1:]), (row_index, values))

    def test_gravimetric_sampling_tables_remain_blank_without_gravimetric_source(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            doc = self._generate(temp_dir)
            # 表22~35是称重法（M0/M1/加样量）表格，Excel 只有光度法吸光度数据，
            # 不能因为同属“加样准确度与重复性检测”章节就跨方法填充。
            for table_index in range(22, 36):
                matrix = physical_table_matrix(doc.tables[table_index])
                for row_index in range(2, 12):
                    values = matrix[row_index]
                    # 序号列保留模板原值，其余测量列必须为空。
                    self.assertTrue(all(value == "" for value in values[1:4]), (table_index, row_index, values))
                    self.assertTrue(all(value == "" for value in values[5:8]), (table_index, row_index, values))
                for row_index, values in enumerate(matrix):
                    if "符合要求" in "".join(values):
                        self.fail((table_index, row_index, values))
                    if len(values) >= 3 and (values[0].strip() == "结论" or values[1].strip() in {"正确度", "重复性"}):
                        self.assertTrue(all(value == "" for value in values[2:]), (table_index, row_index, values))


if __name__ == "__main__":
    unittest.main()
